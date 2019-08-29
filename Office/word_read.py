

# def getScores(path):
#     scores = []
#     document = Document(path)  # 读入文件
#     tables = document.tables  # 获取文件中的表格集
#     table = tables[0]  # 获取文件中的第一个表格
#     for i in range(6, 19):  # 从表格第二行开始循环读取表格数据
#         q = table.cell(i, 4).text
#         a = table.cell(i, 5).text  # + "" +table.cell(i,6).text+table.cell(i,7).text #+ table.cell(i,7).text
#         b = table.cell(i, 6).text
#         c = table.cell(i, 7).text
#         d = table.cell(i, 8).text
#         e = table.cell(i, 9).text
#         # cell(i,0)表示第(i+1)行第1列数据，以此类推
#         res = {'a': a, 'b': b, 'c': c, 'd': d, 'e': e}
#         for k, v in res.items():
#             if v == '1':
#                 score = [q, k]
#                 scores.append(score)
#     return scores
from docx import Document
import pandas as pd
def getScores(path):
    scores = {}
    document = Document(path)  # 读入文件
    tables = document.tables  # 获取文件中的表格集
    table = tables[0]  # 获取文件中的第一个表格
    for i in range(6, 19):  # 从表格第二行开始循环读取表格数据
        q = table.cell(i, 4).text
        a = table.cell(i, 5).text  # + "" +table.cell(i,6).text+table.cell(i,7).text #+ table.cell(i,7).text
        b = table.cell(i, 6).text
        c = table.cell(i, 7).text
        d = table.cell(i, 8).text
        e = table.cell(i, 9).text
        # cell(i,0)表示第(i+1)行第1列数据，以此类推
        res = {'a': a, 'b': b, 'c': c, 'd': d, 'e': e}
        print(res)
        for k, v in res.items():
            if v=='1':
                score = {q:k}
                scores.update(score)
    return scores
p = r'C:\Users\Administrator\Desktop\attachment\2018210168、龙吟、10.20上午、方建平.docx'
print(getScores(p))
# from os_python import getPath
# paths,names = getPath(r'C:\Users\Administrator\Desktop\attachment')
# scores =[]
# for p in paths:
#     try:
#         id = p.strip('C:\\Users\\Administrator\\Desktop\\attachment\\').strip('.docx')
#         s_id = {'id':id}
#         score = getScores(p)
#         score.update(s_id)
#         scores.append(score)
#     except:
#         pass
# df = pd.DataFrame(scores)
# df.to_excel('10月20日.xlsx')







